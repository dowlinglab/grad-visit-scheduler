"""Additional tests for rich SolutionResult/SolutionSet behavior."""

from __future__ import annotations

import os
import warnings
from pathlib import Path

import matplotlib
import matplotlib.pyplot as plt
import pytest
import pyomo.environ as pyo
from pyomo.common.errors import ApplicationError

from grad_visit_scheduler import (
    Scheduler,
    Mode,
    Solver,
    SolutionSet,
    scheduler_from_configs,
    export_visitor_docx,
)
from grad_visit_scheduler.data import load_visitor_csv

matplotlib.use("Agg")


def _solver_available(name: str) -> bool:
    """Return whether a named Pyomo solver backend is available."""
    try:
        if name == "highs":
            return pyo.SolverFactory("appsi_highs").available() or pyo.SolverFactory("highs").available()
        return pyo.SolverFactory(name).available()
    except ApplicationError:
        return False


def _build_solved_scheduler() -> Scheduler:
    """Build and solve the formulation example with HiGHS."""
    root = Path(__file__).parents[1]
    s = scheduler_from_configs(
        root / "examples" / "faculty_formulation.yaml",
        root / "examples" / "config_formulation.yaml",
        root / "examples" / "data_formulation_visitors.csv",
        mode=Mode.NO_OFFSET,
        solver=Solver.HIGHS,
    )
    s.schedule_visitors(
        group_penalty=0.2,
        min_visitors=2,
        max_visitors=8,
        min_faculty=1,
        max_group=2,
        enforce_breaks=True,
        tee=False,
        run_name="unit",
    )
    return s


def _build_tiny_scheduler(tmp_path: Path, mode: Mode = Mode.BUILDING_A_FIRST) -> Scheduler:
    """Build a tiny scheduler for top-N edge-path tests."""
    csv_path = tmp_path / "visitors.csv"
    csv_path.write_text("Name,Prof1,Area1,Area2\nVisitor 1,Faculty A,Area1,Area1\n", encoding="utf-8")

    times_by_building = {
        "BuildingA": ["1:00-1:25"],
        "BuildingB": ["1:00-1:25"],
    }
    faculty_catalog = {
        "Faculty A": {
            "building": "BuildingA",
            "room": "101",
            "areas": ["Area1"],
            "status": "active",
        }
    }
    return Scheduler(
        times_by_building=times_by_building,
        student_data_filename=str(csv_path),
        mode=mode,
        solver=Solver.HIGHS,
        faculty_catalog=faculty_catalog,
    )


def _build_two_slot_single_faculty_scheduler(tmp_path: Path) -> Scheduler:
    """Build a tiny scheduler with exactly two tied-optimal schedules."""
    csv_path = tmp_path / "visitors_two_slot.csv"
    csv_path.write_text("Name,Prof1,Area1,Area2\nVisitor 1,Faculty A,Area1,Area1\n", encoding="utf-8")

    times_by_building = {
        "BuildingA": ["1:00-1:25", "1:30-1:55"],
        "BuildingB": ["1:00-1:25", "1:30-1:55"],
    }
    faculty_catalog = {
        "Faculty A": {
            "building": "BuildingA",
            "room": "101",
            "areas": ["Area1"],
            "status": "active",
        }
    }
    return Scheduler(
        times_by_building=times_by_building,
        student_data_filename=str(csv_path),
        mode=Mode.BUILDING_A_FIRST,
        solver=Solver.HIGHS,
        faculty_catalog=faculty_catalog,
    )


def test_load_visitor_csv(tmp_path: Path):
    """Data helper should read visitor CSV into a DataFrame."""
    path = tmp_path / "v.csv"
    path.write_text("Name,Prof1\nA,X\n", encoding="utf-8")
    df = load_visitor_csv(str(path))
    assert list(df.columns) == ["Name", "Prof1"]
    assert df.iloc[0]["Name"] == "A"


def test_scheduler_data_helpers(tmp_path: Path):
    """Basic data helper methods should mutate scheduler state as expected."""
    s = _build_tiny_scheduler(tmp_path)

    s.add_external_faculty("External Y", building="BuildingB", room="B1", areas=["Area1"], available=[1])
    assert "External Y" in s.faculty
    assert s.faculty["External Y"]["room"] == "B1"

    s.specify_limited_student_availability({"Visitor 1": [1]})
    assert s.students_available["Visitor 1"] == [1]


def test_current_solution_result_raises_before_solve(tmp_path: Path):
    """Accessing current solution snapshot before solve should raise."""
    s = _build_tiny_scheduler(tmp_path)
    with pytest.raises(RuntimeError, match="No feasible solution"):
        s.current_solution()


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_schedule_visitors_returns_solution_result():
    """Single-solve API should return a SolutionResult for feasible solves."""
    s = _build_solved_scheduler()
    sol = s.schedule_visitors(
        group_penalty=0.2,
        min_visitors=2,
        max_visitors=8,
        min_faculty=1,
        max_group=2,
        enforce_breaks=True,
        tee=False,
        run_name="single_return",
    )
    assert sol is not None
    assert sol.rank == 1
    assert sol.objective_value == pytest.approx(float(pyo.value(s.model.obj)))
    assert s.current_solution().active_meetings == sol.active_meetings
    assert s.last_solution_set is None


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_schedule_visitors_returns_none_when_infeasible(tmp_path: Path):
    """Single-solve API should return None when model is infeasible."""
    s = _build_tiny_scheduler(tmp_path)
    sol = s.schedule_visitors(
        group_penalty=0.1,
        min_visitors=2,
        max_visitors=2,
        min_faculty=1,
        max_group=1,
        enforce_breaks=False,
        tee=False,
        run_name="single_infeasible",
    )
    assert sol is None
    assert not s.has_feasible_solution()
    with pytest.raises(RuntimeError, match="No feasible solution"):
        s.current_solution()


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_solution_result_single_solution_api(tmp_path: Path):
    """SolutionResult should support summary, plotting, and DOCX export."""
    pytest.importorskip("docx")
    s = _build_solved_scheduler()
    top = s.schedule_visitors_top_n(
        n_solutions=2,
        group_penalty=0.2,
        min_visitors=2,
        max_visitors=8,
        min_faculty=1,
        max_group=2,
        enforce_breaks=True,
        tee=False,
        run_name="sol_api",
    )
    assert len(top) >= 1

    sol = top.get(1)
    row = sol.summary_row(best_objective=sol.objective_value)
    assert row["rank"] == 1
    assert row["objective_gap_from_best"] == 0.0
    assert row["num_assignments"] == len(sol.active_meetings)
    assert row["num_visitors_scheduled"] >= 1

    any_meeting = next(iter(sol.active_meetings))
    assert sol.meeting_assigned(*any_meeting)

    cwd = Path.cwd()
    try:
        os.chdir(tmp_path)
        sol.plot_visitor_schedule(save_files=False, show_solution_rank=True)
        visitor_title = plt.gcf().axes[0].get_title()
        assert "Solution Rank 1" in visitor_title

        sol.plot_faculty_schedule(save_files=False, show_solution_rank=False)
        faculty_title = plt.gcf().axes[0].get_title()
        assert "Solution Rank" not in faculty_title

        visitor_paths = sol.plot_visitor_schedule(save_files=True, show_solution_rank=True)
        faculty_paths = sol.plot_faculty_schedule(save_files=True, show_solution_rank=True)
        docx_path = sol.export_visitor_docx("single_solution.docx")
    finally:
        os.chdir(cwd)

    assert visitor_paths is not None
    assert faculty_paths is not None
    visitor_png = tmp_path / visitor_paths[0]
    faculty_png = tmp_path / faculty_paths[0]
    docx_file = tmp_path / docx_path.name
    assert visitor_png.exists() and visitor_png.stat().st_size > 0
    assert faculty_png.exists() and faculty_png.stat().st_size > 0
    assert docx_file.exists() and docx_file.stat().st_size > 0

    from docx import Document

    document = Document(str(docx_file))
    assert len(document.tables) == len(sol.visitors)
    assert len(document.tables[0].rows) == len(sol.time_slots)


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_modern_single_solution_workflow_has_no_legacy_futurewarnings(tmp_path: Path):
    """Modern single-solve workflow should avoid legacy wrapper warnings."""
    pytest.importorskip("docx")
    s = _build_solved_scheduler()
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        sol = s.schedule_visitors(
            group_penalty=0.2,
            min_visitors=2,
            max_visitors=8,
            min_faculty=1,
            max_group=2,
            enforce_breaks=True,
            tee=False,
            run_name="modern_single",
        )
        assert sol is not None
        cwd = Path.cwd()
        try:
            os.chdir(tmp_path)
            sol.plot_visitor_schedule(save_files=False, show_solution_rank=False)
            sol.plot_faculty_schedule(save_files=False, show_solution_rank=False)
            sol.export_visitor_docx("modern_single.docx")
        finally:
            os.chdir(cwd)
    legacy_msgs = [
        str(w.message)
        for w in caught
        if issubclass(w.category, FutureWarning) and "legacy" in str(w.message).lower()
    ]
    assert legacy_msgs == []


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_solution_set_summarize_options(tmp_path: Path):
    """Summarize helper should handle optional arguments and docx export."""
    pytest.importorskip("docx")
    s = _build_solved_scheduler()
    top = s.schedule_visitors_top_n(
        n_solutions=2,
        group_penalty=0.2,
        min_visitors=2,
        max_visitors=8,
        min_faculty=1,
        max_group=2,
        enforce_breaks=True,
        tee=False,
        run_name="sum_opts",
    )

    cwd = Path.cwd()
    try:
        os.chdir(tmp_path)
        report = top.summarize(
            compact_columns=["rank", "objective_value", "not_a_column"],
            ranks_to_plot=(1, 99),
            save_files=False,
            export_docx=True,
            docx_prefix="summary_docx",
        )
    finally:
        os.chdir(cwd)

    assert list(report["compact"].columns) == ["rank", "objective_value"]
    assert report["plotted_ranks"] == (1,)
    assert report["visitor_plot_files"] == ()
    assert report["faculty_plot_files"] == ()
    assert len(report["docx_files"]) == len(top)
    assert report["summary"].shape[0] == len(top)
    assert report["summary"]["rank"].tolist() == list(range(1, len(top) + 1))
    for p in report["docx_files"]:
        docx_path = tmp_path / Path(p).name
        assert docx_path.exists()
        assert docx_path.stat().st_size > 0


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_scheduler_legacy_wrappers_emit_futurewarning(tmp_path: Path):
    """Legacy scheduler wrappers should warn and still execute."""
    pytest.importorskip("docx")
    s = _build_solved_scheduler()

    cwd = Path.cwd()
    try:
        os.chdir(tmp_path)
        with pytest.warns(FutureWarning, match="legacy interface"):
            s.show_faculty_schedule(save_files=False)
        with pytest.warns(FutureWarning, match="legacy interface"):
            s.show_visitor_schedule(save_files=False)
        with pytest.warns(FutureWarning, match="legacy interface"):
            out = s.export_visitor_docx("legacy.docx")
    finally:
        os.chdir(cwd)

    legacy_docx = tmp_path / out.name
    assert legacy_docx.exists()
    assert legacy_docx.stat().st_size > 0


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_top_level_export_helper_emits_futurewarning(tmp_path: Path):
    """Top-level export helper should warn and export a DOCX file."""
    pytest.importorskip("docx")
    s = _build_solved_scheduler()

    cwd = Path.cwd()
    try:
        os.chdir(tmp_path)
        with pytest.warns(FutureWarning, match="legacy helper"):
            out = export_visitor_docx(s, "top_level.docx")
    finally:
        os.chdir(cwd)

    top_docx = tmp_path / out.name
    assert top_docx.exists()
    assert top_docx.stat().st_size > 0


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_show_utility_and_plot_preferences():
    """Utility and preference visual helpers should return matplotlib objects."""
    s = _build_solved_scheduler()
    fig1, ax1 = s.plot_preferences()
    assert fig1 is not None
    assert ax1 is not None
    assert ax1.get_xlabel() == "Faculty"
    assert ax1.get_ylabel() == "Prospective Graduate Students"
    assert "Interview Preferences" in ax1.get_title()

    fig2, ax2 = s.show_utility()
    assert fig2 is not None
    assert ax2 is not None
    assert "Total Utility =" in ax2.get_title()
    assert len(ax2.collections) >= 1


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_check_requests_runs(capsys):
    """Request diagnostic should run without raising for solved models."""
    s = _build_solved_scheduler()
    s.check_requests()
    captured = capsys.readouterr()
    lines = [line.strip() for line in captured.out.splitlines() if line.strip()]
    meeting_lines = [line for line in lines if "visitors with" in line]
    assert meeting_lines
    total_visitors_reported = 0
    for line in meeting_lines:
        parts = line.split()
        total_visitors_reported += int(parts[0])
    assert total_visitors_reported == len(s.student_data.index)


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_check_requests_raises_when_infeasible(tmp_path: Path):
    """check_requests should raise RuntimeError on infeasible solve states."""
    s = _build_tiny_scheduler(tmp_path)
    s.schedule_visitors(
        group_penalty=0.1,
        min_visitors=2,
        max_visitors=2,
        min_faculty=1,
        max_group=1,
        enforce_breaks=False,
        tee=False,
        run_name="infeasible_for_requests",
    )
    assert not s.has_feasible_solution()
    with pytest.raises(RuntimeError, match="Termination:"):
        s.check_requests()


def test_top_n_invalid_n_raises(tmp_path: Path):
    """Top-N solve should reject n_solutions < 1."""
    s = _build_tiny_scheduler(tmp_path)
    with pytest.raises(ValueError, match="at least 1"):
        s.schedule_visitors_top_n(n_solutions=0)


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_building_b_first_mode_builds_and_solves(tmp_path: Path):
    """BUILDING_B_FIRST branch should be exercised and remain solvable."""
    csv_path = tmp_path / "visitors_mode_b.csv"
    csv_path.write_text("Name,Prof1,Area1,Area2\nVisitor 1,Faculty A,Area1,Area1\n", encoding="utf-8")
    times_by_building = {
        "BuildingA": ["1:00-1:25", "1:30-1:55"],
        "BuildingB": ["1:00-1:25", "1:30-1:55"],
    }
    faculty_catalog = {
        "Faculty A": {
            "building": "BuildingB",
            "room": "201",
            "areas": ["Area1"],
            "status": "active",
        }
    }
    s = Scheduler(
        times_by_building=times_by_building,
        student_data_filename=str(csv_path),
        mode=Mode.BUILDING_B_FIRST,
        solver=Solver.HIGHS,
        faculty_catalog=faculty_catalog,
    )
    top = s.schedule_visitors_top_n(
        n_solutions=2,
        group_penalty=0.1,
        min_visitors=0,
        max_visitors=2,
        min_faculty=1,
        max_group=1,
        enforce_breaks=False,
        tee=False,
        run_name="mode_b_first",
    )
    assert len(top) >= 1


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_top_n_returns_fewer_than_requested_when_unique(tmp_path: Path):
    """Top-N should stop early when no more unique feasible schedules exist."""
    s = _build_tiny_scheduler(tmp_path)
    top = s.schedule_visitors_top_n(
        n_solutions=5,
        group_penalty=0.1,
        min_visitors=1,
        max_visitors=1,
        min_faculty=1,
        max_group=1,
        enforce_breaks=False,
        tee=False,
        run_name="tiny",
    )
    assert len(top) == 1


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_top_n_first_infeasible_path_and_report(tmp_path: Path):
    """If first top-N solve is infeasible, report state should still be populated."""
    s = _build_tiny_scheduler(tmp_path)
    top = s.schedule_visitors_top_n(
        n_solutions=3,
        group_penalty=0.1,
        min_visitors=2,
        max_visitors=2,
        min_faculty=1,
        max_group=1,
        enforce_breaks=False,
        tee=False,
        run_name="tiny_infeasible",
    )
    assert len(top) == 0
    assert not s.has_feasible_solution()
    report = s.infeasibility_report()
    assert "Termination:" in report
    assert "Infeasibility likely: min_visitors requirement exceeds total capacity." in report
    assert "Faculty with insufficient availability for min_visitors:" in report
    assert "Capacity (students): 1, Capacity (faculty): 1, Effective capacity: 1" in report
    assert "Required meetings from faculty (min_visitors): 2" in report


def test_infeasibility_report_without_results(tmp_path: Path):
    """Infeasibility report should handle pre-solve state."""
    s = _build_tiny_scheduler(tmp_path)
    assert s.infeasibility_report() == "No solver results available."


def test_solution_set_empty_and_rank_bounds():
    """SolutionSet should handle empty collection and invalid rank access."""
    empty = SolutionSet([])
    assert len(empty) == 0
    with pytest.raises(IndexError, match="out of bounds"):
        empty.get(1)

    summary = empty.to_dataframe()
    assert summary.empty

    report = empty.summarize(ranks_to_plot=(1, 2), save_files=False, export_docx=True)
    assert report["summary"].empty
    assert report["compact"].empty
    assert report["plotted_ranks"] == ()
    assert report["visitor_plot_files"] == ()
    assert report["faculty_plot_files"] == ()
    assert report["docx_files"] == ()


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_top_n_no_good_cut_exhausts_tied_two_solution_case(tmp_path: Path):
    """No-good cuts should enumerate exactly two tied slot choices then stop."""
    s = _build_two_slot_single_faculty_scheduler(tmp_path)
    top = s.schedule_visitors_top_n(
        n_solutions=3,
        group_penalty=0.1,
        min_visitors=0,
        max_visitors=1,
        min_faculty=1,
        max_group=1,
        enforce_breaks=False,
        tee=False,
        run_name="two_slot_tied",
    )

    assert len(top) == 2
    assert hasattr(s.model, "no_good_cuts")
    assert len(s.model.no_good_cuts) == 2

    # Each solution has exactly one meeting, and they should differ by slot.
    slot_by_rank = {}
    for sol in top.solutions:
        assert len(sol.active_meetings) == 1
        (_, _, slot) = next(iter(sol.active_meetings))
        slot_by_rank[sol.rank] = slot
    assert set(slot_by_rank.values()) == {1, 2}
    assert top.get(1).objective_value == pytest.approx(top.get(2).objective_value)


@pytest.mark.skipif(not _solver_available("highs"), reason="HiGHS solver unavailable")
def test_top_n_exhaustion_preserves_last_feasible_solver_state(tmp_path: Path):
    """After top-N exhaustion, scheduler state should remain at last feasible solve."""
    s = _build_two_slot_single_faculty_scheduler(tmp_path)
    top = s.schedule_visitors_top_n(
        n_solutions=5,
        group_penalty=0.1,
        min_visitors=0,
        max_visitors=1,
        min_faculty=1,
        max_group=1,
        enforce_breaks=False,
        tee=False,
        run_name="two_slot_state",
    )

    # Model has only two unique schedules; top-N exits on a later infeasible attempt.
    assert len(top) == 2
    assert s.last_solution_set is top

    # Current scheduler-facing state remains feasible and points to the last
    # feasible solution encountered before exhaustion.
    assert s.has_feasible_solution()
    assert s.infeasibility_report() == "Solution is feasible."
    current = s._current_solution_result()
    assert current.rank == 1
    assert current.active_meetings == top.get(len(top)).active_meetings
