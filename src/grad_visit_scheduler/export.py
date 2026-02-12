"""Export helpers."""
from __future__ import annotations

from pathlib import Path


def export_visitor_docx(
    scheduler,
    filename,
    *,
    solution=None,
    building: str | None = None,
    font_name: str = "Arial",
    font_size_pt: int = 11,
    include_breaks: bool = True,
):
    """Export a solved visitor schedule to a DOCX document.

    Parameters
    ----------
    scheduler:
        Solved scheduler instance containing ``model`` and metadata.
    filename:
        Output DOCX filename.
    solution:
        Optional ``SolutionResult`` snapshot to export. If omitted, uses the
        scheduler's most recently solved model assignment.
    building:
        Building key used to pick displayed time labels. Defaults to the first
        configured building.
    font_name:
        Font family used in the generated document.
    font_size_pt:
        Font size in points.
    include_breaks:
        If ``True``, include explicit "Break" rows for unscheduled slots.

    Returns
    -------
    pathlib.Path
        Path to the written DOCX file.

    Raises
    ------
    RuntimeError
        If the scheduler does not have a feasible solution.
    ImportError
        If ``python-docx`` is not installed.
    """
    if solution is None and not scheduler.has_feasible_solution():
        raise RuntimeError(
            f"No feasible solution available (termination: {getattr(scheduler, 'last_termination_condition', None)})."
        )

    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:  # pragma: no cover - environment dependent
        raise ImportError("python-docx is required to export schedules to .docx") from exc

    output_path = Path(filename)

    document = Document()

    def format_font(run):
        run.font.size = Pt(font_size_pt)
        run.font.name = font_name

    if building is None:
        building = next(iter(scheduler.times_by_building))

    times = scheduler.times_by_building[building]
    visitors, faculty, time_slots = scheduler._solution_axes_sets(solution=solution)

    for visitor in visitors:
        p = document.add_paragraph()
        run = p.add_run(visitor)
        format_font(run)

        table = document.add_table(rows=len(time_slots), cols=3)

        for i, t in enumerate(time_slots):
            row = table.rows[i].cells

            tm = times[t - 1]
            start, end = tm.split("-")
            row[0].text = f"{start.strip()} - {end.strip()} pm"

            matched_faculty = [f for f in faculty if scheduler._meeting_assigned(visitor, f, t, solution=solution)]
            if matched_faculty:
                matched = matched_faculty[0]
                bldg = scheduler.faculty[matched]["building"]
                row[1].text = "Prof. " + matched
                row[2].text = scheduler.faculty[matched]["room"] + " " + bldg
            elif include_breaks:
                row[1].text = "Break"
                row[2].text = " "

            for j in range(3):
                if row[j].paragraphs and row[j].paragraphs[0].runs:
                    format_font(row[j].paragraphs[0].runs[0])
                elif row[j].paragraphs:
                    run = row[j].paragraphs[0].add_run("")
                    format_font(run)

        document.add_paragraph(" ")

    document.save(str(output_path))
    return output_path
