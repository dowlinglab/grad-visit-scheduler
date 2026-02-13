"""Core scheduling model, solver integration, and plotting helpers."""

import pyomo.environ as pyo
from pyomo.opt import SolverStatus, TerminationCondition
from pyomo.core import Suffix
from dataclasses import dataclass, replace
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import warnings
import os
from math import isnan
import re
from enum import Enum
from collections import Counter
from pathlib import Path

class Solver(Enum):
    """Supported optimization solver backends."""

    HIGHS = 1
    CBC = 2
    GUROBI = 3
    GUROBI_IIS = 4


class FacultyStatus(Enum):
    """Faculty catalog status labels."""

    ACTIVE = "active"
    LEGACY = "legacy"
    EXTERNAL = "external"

class Mode(Enum):
    """Building sequencing constraints for visitor movement."""

    BUILDING_A_FIRST = 1
    BUILDING_B_FIRST = 2
    NO_OFFSET = 3


class MovementPolicy(Enum):
    """Building movement policy for multi-building schedules."""

    NONE = "none"
    TRAVEL_TIME = "travel_time"

def schedule_axes(figsize,nslots=7):
    """Create a matplotlib axis formatted for visit-day schedule plots.

    Parameters
    ----------
    figsize:
        Figure size passed to ``plt.subplots``.
    nslots:
        Number of schedule slots to display.

    Returns
    -------
    matplotlib.axes.Axes
        Configured axis with time ticks and grid lines.
    """
    fig, ax = plt.subplots(1, 1, figsize=figsize)
    # 30 minute time slots, starting at 1:00 PM, + 10 minutes to
    # draw the last tick
    xticks = [i for i in range(60, 30*nslots + 60 + 10, 15)]
    xlabels = [f"{t//60}:{t%60:02d}" for t in xticks]
    ax.set_xticks(ticks=xticks, labels=xlabels)
    for t in xticks:
        ax.axvline(t, lw=1, alpha=0.3, color='b')
    ax.spines[['left', 'top', 'right', 'bottom']].set_visible(False)
    ax.set_xlabel("Time (PM)")
    return ax
    
def slot2min(slot):
    """Convert a slot string like ``'1:00-1:25'`` to integer minutes.

    Parameters
    ----------
    slot:
        Schedule slot label in ``start-end`` format.

    Returns
    -------
    tuple[int, int]
        Start and end times in minutes.
    """
    start, stop = slot.split("-")
    a, b = start.split(":")
    c, d = stop.split(":")
    return 60*int(a) + int(b), 60*int(c) + int(d)

def abbreviate_name(full_name):
    """Abbreviate a full name for compact schedule labels.

    Examples
    --------
    ``"Jane Doe" -> "Jane D."``
    ``"John Michael Doe" -> "John M. D."``

    Notes
    -----
    If non-initial tokens are not alphabetic (for example ``"Visitor 01"``),
    the original name is returned unchanged to avoid ambiguous labels.
    """
    if not full_name:  # Check for empty or None input
        return ""

    full_name = str(full_name).strip()
    name_parts = full_name.split()

    if not name_parts:  # Check for empty after split
        return ""

    # Avoid collapsing labels like "Visitor 01" into "Visitor 0."
    if any(not any(ch.isalpha() for ch in part) for part in name_parts[1:]):
        return full_name

    num_parts = len(name_parts)

    if num_parts == 1:  # Only first name
        return name_parts[0]
    elif num_parts == 2:  # First and last name
        return f"{name_parts[0]} {name_parts[1][0]}."
    else:  # First, middle, and last names (or more middle names)
        first_name = name_parts[0]
        last_name = name_parts[-1]
        middle_initials = ""

        for part in name_parts[1:-1]:  # Iterate through middle names
            middle_initials += part[0] + ". "

        return f"{first_name} {middle_initials}{last_name[0]}."


DEFAULT_COLOR_CYCLE = [
    "#8ecae6",
    "#90be6d",
    "#f4a261",
    "#e76f51",
    "#bdb2ff",
    "#ffd6a5",
]
BOX_ALPHA = 1.0


@dataclass(frozen=True)
class SolutionContext:
    """Immutable rendering and reporting context shared by solution snapshots."""

    times_by_building: dict[str, tuple[str, ...]]
    faculty: dict[str, dict[str, object]]
    box_colors: dict[str, str]
    number_time_slots: int
    run_name: str
    student_preferences: dict[tuple[str, str], float]
    requests: dict[str, tuple[str, ...]]
    legacy_faculty: frozenset[str]
    external_faculty: frozenset[str]

    def is_legacy(self, faculty_name: str) -> bool:
        """Return whether a faculty entry is marked legacy."""
        return faculty_name in self.legacy_faculty

    def is_external(self, faculty_name: str) -> bool:
        """Return whether a faculty entry is marked external."""
        return faculty_name in self.external_faculty


@dataclass(frozen=True)
class SolutionResult:
    """Rich, self-contained representation of one solved schedule.

    A ``SolutionResult`` is intentionally independent from the mutable
    :class:`Scheduler` and Pyomo model state. It contains all assignment and
    metadata needed to:

    - inspect per-solution quality statistics,
    - produce schedule visualizations,
    - export visitor DOCX schedules,
    - serialize/pickle for later analysis.

    Parameters
    ----------
    rank:
        1-based rank in the top-N solve sequence.
    objective_value:
        Objective value for this ranked solution.
    termination_condition:
        Solver termination label (string form).
    solver_status:
        Solver status label (string form).
    visitors:
        Ordered visitor labels included in the solved model.
    faculty:
        Ordered faculty labels included in the solved model.
    time_slots:
        Ordered integer time-slot indices.
    active_meetings:
        Set of active binary assignments ``(visitor, faculty, slot)``.
    context:
        Immutable rendering/reporting context copied from the scheduler at solve
        time (times, faculty metadata, colors, preferences, requests, etc.).
    """

    rank: int
    objective_value: float
    termination_condition: str
    solver_status: str
    visitors: tuple[str, ...]
    faculty: tuple[str, ...]
    time_slots: tuple[int, ...]
    active_meetings: frozenset[tuple[str, str, int]]
    context: SolutionContext

    def meeting_assigned(self, visitor: str, faculty: str, time_slot: int) -> bool:
        """Return whether the visitor/faculty/time assignment is active."""
        return (visitor, faculty, time_slot) in self.active_meetings

    def _meeting_sizes(self):
        meeting_sizes = {}
        for _, faculty, time_slot in self.active_meetings:
            key = (faculty, time_slot)
            meeting_sizes[key] = meeting_sizes.get(key, 0) + 1
        return meeting_sizes

    def summary_row(self, best_objective=None):
        """Return one summary row dictionary for this solution."""
        visitor_counts = {v: 0 for v in self.visitors}
        faculty_counts = {f: 0 for f in self.faculty}
        requested_meetings = 0
        weighted_preference_sum = 0.0
        legacy_meetings = 0
        external_meetings = 0

        for visitor, faculty, _ in self.active_meetings:
            visitor_counts[visitor] += 1
            faculty_counts[faculty] += 1
            weighted_preference_sum += float(
                self.context.student_preferences.get((visitor, faculty), 0.0)
            )
            if faculty in self.context.requests.get(visitor, ()):
                requested_meetings += 1
            if self.context.is_legacy(faculty):
                legacy_meetings += 1
            if self.context.is_external(faculty):
                external_meetings += 1

        meeting_sizes = self._meeting_sizes()
        visitor_loads = list(visitor_counts.values())
        faculty_loads = list(faculty_counts.values())

        return {
            "rank": self.rank,
            "objective_value": self.objective_value,
            "objective_gap_from_best": 0.0 if best_objective is None else float(best_objective - self.objective_value),
            "termination_condition": self.termination_condition,
            "solver_status": self.solver_status,
            "num_assignments": len(self.active_meetings),
            "num_requested_assignments": requested_meetings,
            "weighted_preference_sum": weighted_preference_sum,
            "num_group_slots": sum(1 for n in meeting_sizes.values() if n > 1),
            "num_one_on_one_slots": sum(1 for n in meeting_sizes.values() if n == 1),
            "max_group_size": max(meeting_sizes.values()) if meeting_sizes else 0,
            "num_visitors_scheduled": sum(1 for n in visitor_loads if n > 0),
            "visitor_meetings_min": min(visitor_loads) if visitor_loads else 0,
            "visitor_meetings_avg": float(np.mean(visitor_loads)) if visitor_loads else 0.0,
            "visitor_meetings_max": max(visitor_loads) if visitor_loads else 0,
            "num_faculty_scheduled": sum(1 for n in faculty_loads if n > 0),
            "faculty_meetings_min": min(faculty_loads) if faculty_loads else 0,
            "faculty_meetings_avg": float(np.mean(faculty_loads)) if faculty_loads else 0.0,
            "faculty_meetings_max": max(faculty_loads) if faculty_loads else 0,
            "legacy_assignments": legacy_meetings,
            "external_assignments": external_meetings,
        }

    def _schedule_filename(self, base: str, include_rank: bool = True, suffix: str = ""):
        name = base
        if self.context.run_name:
            name += "_" + self.context.run_name
        if include_rank:
            name += f"_rank{self.rank}"
        if suffix:
            name += suffix
        return name

    def plot_faculty_schedule(
        self,
        save_files=True,
        abbreviate_student_names=True,
        show_solution_rank=False,
        include_rank_in_filename=True,
    ):
        """Plot schedule grouped by faculty for this single solution."""
        ax = schedule_axes(figsize=(12, 10), nslots=self.context.number_time_slots)
        yticks = [-y for y, _ in enumerate(self.faculty)]
        ylabels = [
            f"{f} {self.context.faculty[f]['building']} ({sum(1 for s in self.visitors for t in self.time_slots if self.meeting_assigned(s, f, t)):0.0f})"
            for f in self.faculty
        ]
        ax.set_yticks(yticks, labels=ylabels)
        for f, label in zip(self.faculty, ax.get_yticklabels()):
            if self.context.is_legacy(f):
                label.set_color("red")
        for y in yticks:
            ax.axhline(y, lw=0.4, alpha=0.3, color="b")

        for y, f in enumerate(self.faculty):
            for t in self.context.faculty[f]["avail"]:
                bldg = self.context.faculty[f]["building"]
                start, stop = slot2min(self.context.times_by_building[bldg][t - 1])
                ax.plot(
                    [start, stop],
                    [-y, -y],
                    lw=20,
                    color=self.context.box_colors.get(bldg, "#cccccc"),
                    alpha=BOX_ALPHA,
                    solid_capstyle="butt",
                )
                students = [s for s in self.visitors if self.meeting_assigned(s, f, t)]
                if abbreviate_student_names:
                    students = [abbreviate_name(s) for s in students]
                if students:
                    ax.text((start + stop) / 2, -y, "\n ".join(students), ha="center", va="center", fontsize=8)

        title = "Schedule by Faculty"
        if show_solution_rank:
            title += f" (Solution Rank {self.rank})"
        ax.set_title(title)
        if save_files:
            name = self._schedule_filename("faculty_schedule", include_rank=include_rank_in_filename)
            plt.savefig(name + ".pdf")
            plt.savefig(name + ".png")
            return (name + ".png", name + ".pdf")
        return None

    def plot_visitor_schedule(
        self,
        save_files=True,
        abbreviate_student_names=True,
        show_solution_rank=False,
        include_rank_in_filename=True,
    ):
        """Plot schedule grouped by visitor for this single solution."""
        ax = schedule_axes(figsize=(12, 10), nslots=self.context.number_time_slots)
        students = [abbreviate_name(s) if abbreviate_student_names else s for s in self.visitors]
        yticks = [-y for y, _ in enumerate(students)]
        ax.set_yticks(yticks, labels=students)
        for y in yticks:
            ax.axhline(y, lw=0.4, alpha=0.3, color="b")

        for y, s in enumerate(self.visitors):
            for t in self.time_slots:
                matched = [f for f in self.faculty if self.meeting_assigned(s, f, t)]
                if not matched:
                    continue
                f = matched[0]
                bldg = self.context.faculty[f]["building"]
                start, stop = slot2min(self.context.times_by_building[bldg][t - 1])
                ax.plot(
                    [start, stop],
                    [-y, -y],
                    lw=20,
                    color=self.context.box_colors.get(bldg, "#cccccc"),
                    alpha=BOX_ALPHA,
                    solid_capstyle="butt",
                )
                text_color = "red" if self.context.is_legacy(f) else "black"
                ax.text((start + stop) / 2, -y, f"{f} ({bldg})", ha="center", va="center", fontsize=8, color=text_color)

        title = "Schedule by Visitors"
        if show_solution_rank:
            title += f" (Solution Rank {self.rank})"
        ax.set_title(title)
        if save_files:
            name = self._schedule_filename("visitor_schedule", include_rank=include_rank_in_filename)
            plt.savefig(name + ".pdf")
            plt.savefig(name + ".png")
            return (name + ".png", name + ".pdf")
        return None

    def export_visitor_docx(
        self,
        filename,
        *,
        building: str | None = None,
        font_name: str = "Arial",
        font_size_pt: int = 11,
        include_breaks: bool = True,
    ):
        """Export this solution to a visitor schedule DOCX file."""
        try:
            from docx import Document
            from docx.shared import Pt
        except Exception as exc:  # pragma: no cover - environment dependent
            raise ImportError("python-docx is required to export schedules to .docx") from exc

        output_path = Path(filename)
        document = Document()

        def format_font(run):
            run.font.size = Pt(font_size_pt)
            run.font.name = font_name

        if building is None:
            building = next(iter(self.context.times_by_building))
        times = self.context.times_by_building[building]

        for visitor in self.visitors:
            p = document.add_paragraph()
            run = p.add_run(visitor)
            format_font(run)

            table = document.add_table(rows=len(self.time_slots), cols=3)
            for i, t in enumerate(self.time_slots):
                row = table.rows[i].cells
                start, end = times[t - 1].split("-")
                row[0].text = f"{start.strip()} - {end.strip()} pm"

                matched = [f for f in self.faculty if self.meeting_assigned(visitor, f, t)]
                if matched:
                    f = matched[0]
                    bldg = self.context.faculty[f]["building"]
                    row[1].text = "Prof. " + f
                    row[2].text = self.context.faculty[f]["room"] + " " + bldg
                elif include_breaks:
                    row[1].text = "Break"
                    row[2].text = " "

                for j in range(3):
                    if row[j].paragraphs and row[j].paragraphs[0].runs:
                        format_font(row[j].paragraphs[0].runs[0])
                    elif row[j].paragraphs:
                        run = row[j].paragraphs[0].add_run("")
                        format_font(run)

            document.add_paragraph(" ")

        document.save(str(output_path))
        return output_path


class SolutionSet:
    """Collection object for ranked :class:`SolutionResult` instances."""

    def __init__(self, solutions):
        self.solutions = tuple(solutions)

    def __len__(self):
        return len(self.solutions)

    def best(self):
        """Return the top-ranked solution, if present."""
        return self.get(1)

    def get(self, rank: int):
        """Return a ranked solution (1-based rank)."""
        if rank < 1 or rank > len(self.solutions):
            raise IndexError(f"Rank {rank} is out of bounds for {len(self.solutions)} solutions.")
        return self.solutions[rank - 1]

    def to_dataframe(self):
        """Return a summary dataframe of ranked solution quality statistics."""
        best_obj = self.solutions[0].objective_value if self.solutions else None
        return pd.DataFrame([s.summary_row(best_objective=best_obj) for s in self.solutions])

    def summarize(
        self,
        *,
        compact_columns=None,
        ranks_to_plot=(1, 2),
        save_files=True,
        show_solution_rank=True,
        plot_prefix=None,
        export_docx=False,
        docx_prefix="visitor_schedule_top",
    ):
        """Build a reusable top-N solution review bundle.

        This helper packages the common "inspect top-N solutions" workflow used
        in examples and notebooks:

        1. Build a full summary table with one row per ranked solution.
        2. Build a compact summary view with key comparison columns.
        3. Optionally generate visitor/faculty schedule plots for selected ranks.
        4. Optionally export all ranked solutions to DOCX.

        Parameters
        ----------
        compact_columns:
            Optional list of column names to include in a compact comparison
            table. If omitted, a default set of high-signal columns is used.
        ranks_to_plot:
            Iterable of 1-based ranks to render (for example ``(1, 2)``).
            Invalid ranks are ignored.
        save_files:
            If ``True``, plotting helpers write files to disk and return their
            paths in the output dictionary.
        show_solution_rank:
            If ``True``, append ``"(Solution Rank X)"`` to plot titles for
            generated ranked-solution figures.
        plot_prefix:
            Optional run-name override used while generating plots. If provided,
            output plot filenames become ``*_plot_prefix_rankX.*``.
        export_docx:
            If ``True``, export all ranked solutions to DOCX files.
        docx_prefix:
            Prefix used when ``export_docx=True``. Filenames are generated as
            ``{docx_prefix}_rank{rank}.docx``.

        Returns
        -------
        dict
            Dictionary with the following keys:

            - ``summary``: full summary dataframe from :meth:`to_dataframe`.
            - ``compact``: compact comparison dataframe.
            - ``plotted_ranks``: tuple of ranks that were plotted.
            - ``visitor_plot_files``: tuple of PNG filenames (if saved).
            - ``faculty_plot_files``: tuple of PNG filenames (if saved).
            - ``docx_files``: tuple of exported DOCX filenames.

        Notes
        -----
        - This helper is intentionally non-destructive: it restores the
          scheduler run name after temporary overrides.
        - Plot filenames are returned only when ``save_files=True``.
        - Compact columns that are not present in the summary dataframe are
          silently skipped.
        """
        summary = self.to_dataframe()
        default_compact_columns = [
            "rank",
            "objective_value",
            "objective_gap_from_best",
            "num_assignments",
            "num_group_slots",
            "visitor_meetings_avg",
            "faculty_meetings_avg",
        ]
        selected_columns = compact_columns if compact_columns is not None else default_compact_columns
        compact_cols_present = [c for c in selected_columns if c in summary.columns]
        compact = summary[compact_cols_present].copy()

        visitor_plot_files = []
        faculty_plot_files = []
        plotted_ranks = []

        for rank in ranks_to_plot:
            if rank < 1 or rank > len(self.solutions):
                continue
            plotted_ranks.append(rank)
            solution = self.get(rank)
            if plot_prefix is not None:
                solution = replace(solution, context=replace(solution.context, run_name=plot_prefix))
            if save_files:
                visitor_paths = solution.plot_visitor_schedule(
                    save_files=True, show_solution_rank=show_solution_rank
                )
                faculty_paths = solution.plot_faculty_schedule(
                    save_files=True, show_solution_rank=show_solution_rank
                )
                if visitor_paths:
                    visitor_plot_files.append(visitor_paths[0])
                if faculty_paths:
                    faculty_plot_files.append(faculty_paths[0])
            else:
                solution.plot_visitor_schedule(save_files=False, show_solution_rank=show_solution_rank)
                solution.plot_faculty_schedule(save_files=False, show_solution_rank=show_solution_rank)

        docx_files = []
        if export_docx:
            for p in self.export_visitor_docx_all(prefix=docx_prefix):
                docx_files.append(str(p))

        return {
            "summary": summary,
            "compact": compact,
            "plotted_ranks": tuple(plotted_ranks),
            "visitor_plot_files": tuple(visitor_plot_files),
            "faculty_plot_files": tuple(faculty_plot_files),
            "docx_files": tuple(docx_files),
        }

    def plot_faculty_schedule(self, rank=1, show_solution_rank=True, **kwargs):
        """Plot faculty schedule for the selected ranked solution."""
        return self.get(rank).plot_faculty_schedule(show_solution_rank=show_solution_rank, **kwargs)

    def plot_visitor_schedule(self, rank=1, show_solution_rank=True, **kwargs):
        """Plot visitor schedule for the selected ranked solution."""
        return self.get(rank).plot_visitor_schedule(show_solution_rank=show_solution_rank, **kwargs)

    def export_visitor_docx(self, filename, rank=1, **kwargs):
        """Export a selected ranked solution to DOCX."""
        return self.get(rank).export_visitor_docx(filename, **kwargs)

    def export_visitor_docx_all(self, prefix="visitor_schedule", suffix=".docx", **kwargs):
        """Export all ranked solutions to separate DOCX files."""
        output_paths = []
        for solution in self.solutions:
            filename = Path(f"{prefix}_rank{solution.rank}{suffix}")
            output_paths.append(solution.export_visitor_docx(filename, **kwargs))
        return output_paths


class Scheduler:
    """Scheduler for assigning visitor-faculty meetings across time slots."""
    def __init__(
        self,
        times_by_building,
        student_data_filename,
        mode=None,
        movement=None,
        solver=Solver.HIGHS,
        include_legacy_faculty=False,
        faculty_catalog=None,
        faculty_aliases=None,
    ):
        """Initialize a scheduler with buildings, visitors, and faculty data.

        Parameters
        ----------
        times_by_building:
            Mapping of building names to ordered time-slot labels and optional
            ``"breaks"`` list.
        student_data_filename:
            CSV file with visitor names and ranked preferences.
        mode:
            Legacy building sequencing mode. Prefer ``movement``.
        movement:
            Optional movement configuration dictionary with keys:
            ``policy`` (``"none"`` or ``"travel_time"``),
            ``phase_slot`` (per-building earliest slot), and optional
            ``travel_slots`` (pairwise slot lags for travel-time policies).
        solver:
            Solver backend used by :meth:`schedule_visitors`.
        include_legacy_faculty:
            If ``True``, include all legacy faculty entries from the catalog.
        faculty_catalog:
            Optional faculty catalog dictionary. If omitted, a small default
            synthetic catalog is created.
        faculty_aliases:
            Optional mapping of alias name to canonical faculty name.
        """

        # Faculty fields (adjust this depending on number of supplied preferences)
        # Support up to 5 ranked faculty preferences
        self.faculty_fields = [f"Prof{i}" for i in range(1, 6)]

        # Area fields (adjust this depending on number of supplied preferences)
        self.area_fields = ["Area1","Area2"]

        # Process time data
        self._set_time_data(times_by_building)
        self._configure_movement(mode=mode, movement=movement)
        self.box_colors = self._build_box_colors()

        self.external_faculty = {}
        self.faculty_aliases = faculty_aliases or {}
        # Load faculty data
        if faculty_catalog is None:
            self._load_default_faculty_data()
        else:
            self._load_faculty_catalog(faculty_catalog)
        self.include_legacy_faculty = include_legacy_faculty

        # Load student preferences csv file
        self._load_student_preferences(student_data_filename)

        # Create weights using defaults
        self.update_weights()

        # Save default solver and legacy mode marker
        self.mode = mode
        self.solver = solver

    def _set_time_data(self, times_by_building):
        """Validate building slot data and populate scheduling time metadata.

        Parameters
        ----------
        times_by_building:
            Mapping of building names to ordered slot labels, plus optional
            ``"breaks"`` key containing slot indices.
        """

        self.break_times = []
        times_by_building_copy = {}
        for i, (k, v) in enumerate(times_by_building.items()):
            if k != "breaks":
                times_by_building_copy[k] = v
            else:
                self.break_times = v

        self.buildings = list(times_by_building_copy.keys())
        if len(self.buildings) < 1:
            raise ValueError("Run config must define at least one building.")
        self.building_a = self.buildings[0]
        self.building_b = self.buildings[1] if len(self.buildings) > 1 else self.buildings[0]
        self.number_time_slots = len(times_by_building_copy[self.buildings[0]])
        self.time_slots = [i for i in range(1, self.number_time_slots+1) ]

        if len(self.break_times) > 0:
            for t in self.break_times:
                if t not in self.time_slots:
                    raise ValueError(t,"is not a valid break time")

        for i, k in enumerate(times_by_building_copy):
            if len(times_by_building_copy[k]) != self.number_time_slots:
                raise ValueError("Each building should have the same number of timeslots")

        self.times_by_building = times_by_building_copy

    def _configure_movement(self, mode, movement):
        """Normalize movement configuration with legacy mode compatibility."""
        if movement is None:
            movement = {}

        if mode is not None and movement:
            warnings.warn(
                "Both legacy `mode` and `movement` were provided. "
                "Ignoring `mode` and using `movement`.",
                FutureWarning,
                stacklevel=2,
            )

        self.require_break_constraints_default = False
        use_legacy_mode = bool(mode is not None and not movement)
        if use_legacy_mode:
            warnings.warn(
                "Scheduler(mode=...) is a legacy interface. Prefer movement "
                "configuration via run config or Scheduler(..., movement=...).",
                FutureWarning,
                stacklevel=2,
            )

        if use_legacy_mode:
            if len(self.buildings) != 2 and mode in {Mode.BUILDING_A_FIRST, Mode.BUILDING_B_FIRST}:
                raise ValueError("Legacy BUILDING_A_FIRST/BUILDING_B_FIRST modes require exactly two buildings.")
            if mode is Mode.BUILDING_A_FIRST:
                second_slot = min(2, self.number_time_slots)
                movement = {
                    "policy": MovementPolicy.NONE.value,
                    "phase_slot": {self.building_a: 1, self.building_b: second_slot},
                }
            elif mode is Mode.BUILDING_B_FIRST:
                second_slot = min(2, self.number_time_slots)
                movement = {
                    "policy": MovementPolicy.NONE.value,
                    "phase_slot": {self.building_a: second_slot, self.building_b: 1},
                }
            else:
                travel = {b: {bb: (0 if b == bb else 1) for bb in self.buildings} for b in self.buildings}
                movement = {
                    "policy": MovementPolicy.TRAVEL_TIME.value,
                    "phase_slot": {b: 1 for b in self.buildings},
                    "travel_slots": travel,
                }
                # Preserve historical behavior where NO_OFFSET implied explicit breaks.
                self.require_break_constraints_default = True

        policy = str(movement.get("policy", MovementPolicy.NONE.value)).lower()
        if policy not in {MovementPolicy.NONE.value, MovementPolicy.TRAVEL_TIME.value}:
            raise ValueError(f"Unsupported movement policy '{policy}'.")
        self.movement_policy = policy

        phase_slot = dict(movement.get("phase_slot", {}))
        self.building_phase_slot = {}
        for b in self.buildings:
            val = int(phase_slot.get(b, 1))
            if val < 1 or val > self.number_time_slots:
                raise ValueError(
                    f"movement.phase_slot[{b}]={val} is outside valid slot range 1..{self.number_time_slots}."
                )
            self.building_phase_slot[b] = val

        travel_slots = movement.get("travel_slots")
        if self.movement_policy == MovementPolicy.TRAVEL_TIME.value:
            if travel_slots is None:
                travel_slots = {b: {bb: (0 if b == bb else 1) for bb in self.buildings} for b in self.buildings}
            self.travel_slots = {}
            for b_from in self.buildings:
                row = travel_slots.get(b_from)
                if row is None:
                    raise ValueError(f"movement.travel_slots is missing row for '{b_from}'.")
                self.travel_slots[b_from] = {}
                for b_to in self.buildings:
                    if b_to not in row:
                        raise ValueError(
                            f"movement.travel_slots['{b_from}'] is missing destination '{b_to}'."
                        )
                    lag = int(row[b_to])
                    if lag < 0:
                        raise ValueError("movement.travel_slots values must be nonnegative integers.")
                    self.travel_slots[b_from][b_to] = lag
        else:
            self.travel_slots = {b: {bb: 0 for bb in self.buildings} for b in self.buildings}

    def _build_box_colors(self):
        """Return a color map for configured buildings."""
        return {
            bldg: DEFAULT_COLOR_CYCLE[i % len(DEFAULT_COLOR_CYCLE)]
            for i, bldg in enumerate(self.buildings)
        }

    def _load_student_preferences(self, filename):
        """ Load datafile with student preferences
        
        """

        student_data = pd.read_csv(filename)
        if "Name" not in student_data.columns:
            raise ValueError("Student preferences CSV must include a 'Name' column.")
        # Detect which faculty preference columns are present (supports 4 or 5)
        self.faculty_fields = [f for f in self.faculty_fields if f in student_data.columns]

        #for index, row in student_data.iterrows():
        #    print(row)

        #student_data.head()
        if student_data["Name"].duplicated().any():
            dupes = student_data.loc[student_data["Name"].duplicated(), "Name"].tolist()
            raise ValueError(f"Duplicate visitor names found: {dupes}")
        self.student_data = student_data.set_index('Name').sort_index()
        #self.student_data = student_data.sort_index()

        # Loop over rows
        for index, row in self.student_data.iterrows():
            # Loop over faculty columns
            for f in self.faculty_fields:
                if f in self.student_data.columns and pd.notna(row[f]):
                    name = str(row[f]).strip()
                    # Graceful handling of empty / nan-like strings
                    if not name or name.lower() in {"nan", "none", "na"}:
                        continue
                    # Drop faculty first name (after comma)
                    name = name.split(",")[0].strip()
                    # Apply aliases
                    name = self.faculty_aliases.get(name, name)
                    if name:
                        # Drop faculty first name (after comma)
                        row[f] = name

        # Add legacy faculty if they appear in preferences (backwards compatibility)
        self._add_legacy_faculty_from_preferences(include_all=self.include_legacy_faculty)
        # Add external faculty if they appear in preferences
        self._add_external_faculty_from_preferences()

        # To start, no visitors have limited availability
        self.specify_limited_student_availability({})

    def _load_default_faculty_data(self):
        """Load a minimal synthetic faculty catalog for default/demo usage."""
        self.legacy_faculty = {}
        self.faculty = {}
        self.external_faculty = {}

        for i, bldg in enumerate(self.buildings):
            name = f"Faculty {chr(65 + i)}"
            self.faculty[name] = {
                "building": bldg,
                "avail": self.time_slots,
                "areas": [f"Area{i + 1}"],
                "room": f"{100 + i}",
            }

    def _load_faculty_catalog(self, catalog):
        """Load faculty entries from a user-provided catalog dictionary.

        Parameters
        ----------
        catalog:
            Mapping from faculty name to metadata fields.
        """
        self.legacy_faculty = {}
        self.faculty = {}
        self.external_faculty = {}

        for name, info in catalog.items():
            status = str(info.get("status", FacultyStatus.ACTIVE.value)).lower()
            if status not in {FacultyStatus.ACTIVE.value, FacultyStatus.LEGACY.value, FacultyStatus.EXTERNAL.value}:
                raise ValueError(f"Invalid faculty status '{status}' for {name}.")
            entry = {
                "building": info.get("building", self.building_a),
                "avail": self.time_slots,
                "areas": info.get("areas", []),
                "room": info.get("room", ""),
            }
            if status == FacultyStatus.LEGACY.value:
                self.legacy_faculty[name] = entry
            elif status == FacultyStatus.EXTERNAL.value:
                # External faculty default to unavailable unless specified later
                entry["avail"] = []
                self.external_faculty[name] = entry
            else:
                self.faculty[name] = entry

        # Merge external faculty into active faculty
        for name, info in self.external_faculty.items():
            if name not in self.faculty:
                self.faculty[name] = dict(info)

    def _add_legacy_faculty_from_preferences(self, include_all=False):
        """Merge legacy faculty into active scheduling set.

        Parameters
        ----------
        include_all:
            If ``True``, include all legacy faculty regardless of requests.
        """
        if not hasattr(self, "legacy_faculty"):
            return
        legacy_names = set(self.legacy_faculty.keys())
        if include_all:
            for name in legacy_names:
                if name not in self.faculty:
                    self.faculty[name] = dict(self.legacy_faculty[name])
            return
        names_in_prefs = set()
        for f in self.faculty_fields:
            if f in self.student_data.columns:
                names_in_prefs.update(self.student_data[f].dropna().astype(str).str.split(",").str[0].str.strip())
        for name in names_in_prefs:
            if name in legacy_names and name not in self.faculty:
                self.faculty[name] = dict(self.legacy_faculty[name])

    def _add_external_faculty_from_preferences(self):
        """Create placeholder external faculty records from visitor requests."""
        names_in_prefs = set()
        for f in self.faculty_fields:
            if f in self.student_data.columns:
                names_in_prefs.update(self.student_data[f].dropna().astype(str).str.split(",").str[0].str.strip())
        for name in names_in_prefs:
            name = self.faculty_aliases.get(name, name)
            if name and name not in self.faculty and name not in self.external_faculty and not self._is_legacy_faculty(name):
                # Default to the first building with no availability; override via add_external_faculty if needed
                self.external_faculty[name] = {
                    'building': self.building_a,
                    'avail': [],
                    'areas': [],
                    'room': ''
                }
        # Merge external faculty into active faculty
        for name, info in self.external_faculty.items():
            if name not in self.faculty:
                self.faculty[name] = dict(info)

    def add_external_faculty(self, name, building=None, room='', areas=None, available=None):
        """Add or update an external faculty entry.

        Parameters
        ----------
        name:
            Faculty display name.
        building:
            Building key. Defaults to building A.
        room:
            Room label used in exports and plots.
        areas:
            Optional list of research areas for area-based preference boosts.
        available:
            Optional list of available time-slot indices.
        """
        if areas is None:
            areas = []
        if available is None:
            available = self.time_slots
        if building is None:
            building = self.building_a
        self.external_faculty[name] = {
            'building': building,
            'avail': available,
            'areas': areas,
            'room': room
        }
        if name not in self.faculty:
            self.faculty[name] = dict(self.external_faculty[name])

    def _is_legacy_faculty(self, name):
        """Return whether ``name`` belongs to the legacy faculty pool."""
        return hasattr(self, "legacy_faculty") and name in self.legacy_faculty

    def faculty_limited_availability(self, name, available):
        """Set the available time slots for a faculty member.

        Parameters
        ----------
        name:
            Faculty name.
        available:
            List of integer slot indices where meetings are allowed.
        """
        if name not in self.faculty.keys():
            raise ValueError(name,"is not a faculty member... check spelling")

        for t in available:
            if t not in self.time_slots:
                raise ValueError(t,"is not a valid time slot")

        self.faculty[name]['avail'] = available

    def update_weights(self, 
                       faculty_weight = {'Prof1': 4.0, 'Prof2': 3.0, 'Prof3': 2.0, 'Prof4': 1.0, 'Prof5': 0.5},
                       area_weight = {'Area1': 1.0, 'Area2': 0.5},
                       base_weight = 0.2):
        """Set preference weights and rebuild visitor-faculty utility scores.

        Parameters
        ----------
        faculty_weight:
            Either a scalar applied to all ranked faculty columns, or a dict
            keyed by preference column names (for example ``Prof1``).
        area_weight:
            Either a scalar applied to all area columns, or a dict keyed by area
            column names (for example ``Area1``).
        base_weight:
            Baseline utility assigned to any feasible visitor-faculty match.
        """

        if isinstance(faculty_weight, dict):
            self.faculty_weights = {k: v for k, v in faculty_weight.items() if k in self.faculty_fields}
        elif isinstance(faculty_weight, float) or isinstance(faculty_weight, int):
            self.faculty_weights = {}
            for f in self.faculty_fields:
                self.faculty_weights[f] = faculty_weight
        else:
            raise ValueError("faculty_weight must be a float, int, or dict")

        if isinstance(area_weight, dict):
            self.area_weights = area_weight
        elif isinstance(area_weight, float) or isinstance(area_weight, int):
            self.area_weights = {}
            for a in self.area_fields:
                self.area_weights[a] = area_weight
        else:
            raise ValueError("area_weight must be a float, int, or dict")        

        self.base_weight = base_weight

        self._update_student_preferences() 

    def _update_student_preferences(self):
        """Recompute utility weights for each visitor-faculty pair."""

        # Store specific requests by student name
        self.requests = {s:[] for s in self.student_data.index}

        # Initialize dictionary containing preferences
        self.student_preferences = {(s, f): self.base_weight for s in self.student_data.index for f in self.faculty.keys()}

        # Loop over fields Prof1, Prof2, ...
        for i, (k, v) in enumerate(self.faculty_weights.items()):
            # Loop over all students extracting faculty choice for ProfX
            for j, (s, f) in enumerate(self.student_data[k].to_dict().items()):
                if pd.isna(f):
                    continue
                name = str(f).strip()
                if not name or name.lower() in {"nan", "none", "na"}:
                    continue
                # Apply aliases just in case
                name = self.faculty_aliases.get(name, name)
                if name in self.faculty:
                    self.student_preferences[s, name] = v
                    self.requests[s].append(name)
                else:
                    print(f"Unknown faculty preference for {s}: '{name}'. Check spelling or add as external.")
        
        # Loop over area fields, Area1, Area2, ...
        for i, (af, aw) in enumerate(self.area_weights.items()):

            # Loop over all students extracting choice for AreaX
            for j, (s, a) in enumerate(self.student_data[af].to_dict().items()):
                if pd.isna(a):
                    continue
                # Loop over all faculty
                for f in self.faculty.keys():
                    
                    # Check if faculty is in the area
                    if a in self.faculty[f]["areas"]:
                        self.student_preferences[(s, f)] += aw



    def specify_limited_student_availability(self, students_available):
        """Restrict availability for a subset of visitors.

        Parameters
        ----------
        students_available:
            Mapping of visitor name to allowed slot indices. Visitors not listed
            remain available for all time slots.
        """

        students_available

        # Error checking
        for i, (k,v) in enumerate(students_available.items()):
            if k not in self.student_data.index:
                raise ValueError(k,"is not a valid student name")
                for j in v:
                    if j not in self.time_slots:
                        raise ValueError(j,"is not a valid time slot")

        # Assume all other students are available for all time slots
        for s, row in self.student_data.iterrows():
            if s not in students_available.keys():
                students_available[s] = self.time_slots

        # Save in object
        self.students_available = students_available

    def plot_preferences(self):
        """Plot a heatmap of visitor-faculty utility weights.

        Returns
        -------
        tuple
            ``(fig, ax)`` matplotlib objects for further customization.
        """
        fig, ax = plt.subplots(1, 1, figsize=(10, 10))

        df = pd.DataFrame()
        for key, val in self.student_preferences.items():
            s, f = key
            df.loc[s, f] = val

        ax = sns.heatmap(df, cmap="crest", annot=True, square=True, cbar=False)
        ax.set_ylabel('Prospective Graduate Students')
        ax.set_xlabel('Faculty')
        ax.set_title('Graduate Student Interview Preferences')
        return fig, ax

    def schedule_visitors(self, group_penalty=0.1, min_visitors=0, max_visitors=8, min_faculty=0, max_group=2, enforce_breaks=False, tee=False, run_name=''):
        """Solve the mixed-integer scheduling model.

        Parameters
        ----------
        group_penalty:
            Utility penalty for adding a second (or later) visitor to the same
            faculty-time meeting.
        min_visitors:
            Minimum total meetings per available faculty member.
        max_visitors:
            Maximum total meetings per faculty member.
        min_faculty:
            Minimum number of meetings required per visitor.
        max_group:
            Maximum number of visitors allowed in a single faculty-time meeting.
        enforce_breaks:
            If ``True``, enforce explicit break constraints regardless of
            movement policy.
        tee:
            If ``True``, stream solver output.
        run_name:
            Optional suffix for saved plot filenames.
        """

        self.last_solve_params = {
            "group_penalty": group_penalty,
            "min_visitors": min_visitors,
            "max_visitors": max_visitors,
            "min_faculty": min_faculty,
            "max_group": max_group,
            "enforce_breaks": enforce_breaks,
        }
        self._build_model(group_penalty, min_visitors, max_visitors, min_faculty, max_group, enforce_breaks)
        self._solve_model(tee)
        self.run_name = run_name
        self.last_solution_set = None

    def schedule_visitors_top_n(
        self,
        n_solutions=5,
        group_penalty=0.1,
        min_visitors=0,
        max_visitors=8,
        min_faculty=0,
        max_group=2,
        enforce_breaks=False,
        tee=False,
        run_name='',
    ):
        """Solve for up to the best ``n_solutions`` unique schedules.

        Uses no-good integer cuts over the full assignment vector ``y[s, f, t]``
        so each returned solution differs from all previous solutions.
        """
        if n_solutions < 1:
            raise ValueError("n_solutions must be at least 1")

        self.last_solve_params = {
            "group_penalty": group_penalty,
            "min_visitors": min_visitors,
            "max_visitors": max_visitors,
            "min_faculty": min_faculty,
            "max_group": max_group,
            "enforce_breaks": enforce_breaks,
            "n_solutions": n_solutions,
        }
        self._build_model(group_penalty, min_visitors, max_visitors, min_faculty, max_group, enforce_breaks)
        self.run_name = run_name

        solutions = []
        for rank in range(1, n_solutions + 1):
            results = self._solve_model(tee, record_results=False)
            termination = results.solver.termination_condition
            status = results.solver.status

            if termination not in [TerminationCondition.optimal, TerminationCondition.feasible]:
                if not solutions:
                    self.results = results
                    self.last_termination_condition = termination
                    self.last_solver_status = status
                break

            self.results = results
            self.last_termination_condition = termination
            self.last_solver_status = status

            solution = self._snapshot_solution(rank=rank)
            solutions.append(solution)
            self._add_no_good_cut(solution)

        solution_set = SolutionSet(solutions)
        self.last_solution_set = solution_set
        return solution_set

    def _build_model(self, group_penalty, min_visitors, max_visitors, min_faculty, max_group, enforce_breaks=False):
        """Build the Pyomo MILP model for the current scheduler state.

        Parameters
        ----------
        group_penalty:
            Penalty applied to group meetings.
        min_visitors:
            Minimum required meetings per available faculty member.
        max_visitors:
            Maximum total meetings per faculty member.
        min_faculty:
            Minimum required meetings per visitor.
        max_group:
            Maximum number of visitors in a single faculty-time meeting.
        enforce_breaks:
            If ``True``, enforce break constraints during configured break slots.
        """
        m = pyo.ConcreteModel()

        # SETS
        
        m.faculty_available = []
        for f in self.faculty.keys():
            if len(self.faculty[f]["avail"]) > 0:
                m.faculty_available.append(f)
        
        m.visitors = pyo.Set(initialize = self.student_data.index)
        m.faculty = pyo.Set(initialize = m.faculty_available)
        m.time = pyo.RangeSet(1, self.number_time_slots)
        m.buildings = pyo.Set(initialize=self.buildings)
        m.faculty_by_building = pyo.Set(
            m.buildings,
            initialize=lambda m, b: [f for f in m.faculty if self.faculty[f]["building"] == b],
        )
        
        # PARAMETERS
        
        # Penalty for meetings with more than one student at once
        m.penalty = pyo.Param(initialize=group_penalty)

        # student preferences
        @m.Param(m.visitors, m.faculty)
        def weights(m, s, f):
            return self.student_preferences[(s, f)]

        ## DECISION VARIABLES
        
        # y[s, f, t] == 1  <=>  visitor s meets faculty f at time t
        m.y = pyo.Var(m.visitors, m.faculty, m.time, domain=pyo.Binary)
        
        # beyond_one_visitor == number of visitors in excess of one by faculty f at time t
        m.beyond_one_visitor = pyo.Var(m.faculty, m.time, domain=pyo.PositiveReals)
        
        # facutly_too_many_meetings == 1  <=>  faculty f meets more than max_visitors - 2 students
        m.faculty_too_many_meetings = pyo.Var(m.faculty, domain=pyo.Binary)

        if self.movement_policy == MovementPolicy.TRAVEL_TIME.value:
            # in_building[s, b, t] == 1  <=>  visitor s is in building b at time t
            m.in_building = pyo.Var(m.visitors, m.buildings, m.time, domain=pyo.Binary)

        if self.require_break_constraints_default or enforce_breaks:

            if len(self.break_times) == 0:
                raise ValueError("Must specify some break times!")
            
            m.break_options = self.break_times

            m.faculty_breaks = pyo.Var(m.faculty, m.break_options, domain=pyo.Binary)

        # EXPRESSIONS AND OBJECTIVE
        
        # utility is the total value of student preferences that are be accomodated in the schedule
        @m.Expression(m.visitors)
        def utility(m, s):
            return sum(m.weights[s, f] * m.y[s, f, t] for t in m.time for f in m.faculty)

        # assign a penalty if faculty meet more than one student at a time
        @m.Expression(m.faculty)
        def cost(m, f):
            return sum(m.beyond_one_visitor[f, t] for t in m.time)

        @m.Objective(sense=pyo.maximize)
        def obj(m):
            return sum(m.utility[s] for s in m.visitors) - m.penalty * sum(m.cost[f] for f in m.faculty) - 3*m.penalty * sum(m.faculty_too_many_meetings[f] for f in m.faculty)
        
        # CONSTRAINTS

        # no meeting is possible if a faculty member is unavailable
        @m.Constraint(m.visitors, m.faculty, m.time)
        def availability(m, s, f, t):
            return m.y[s, f, t] <= (1 if t in self.faculty[f]['avail'] else 0)
        
        @m.Constraint(m.faculty)
        def min_visitors_constraint(m, f):
            return sum(m.y[s, f, t] for s in m.visitors for t in m.time) >= min_visitors
        
        @m.Constraint(m.faculty)
        def max_visitors_constraint(m, f):
            return sum(m.y[s, f, t] for s in m.visitors for t in m.time) <= max_visitors
    
        # no student can be in more than one meeting at time
        @m.Constraint(m.visitors, m.time)
        def no_simultanous_meetings(m, s, t):
            return sum(m.y[s, f, t] for f in m.faculty) <= 1

        # all students meet a minimum number of faculty
        @m.Constraint(m.visitors)
        def min_faculty(m, s):
            return sum(m.y[s, f, t] for f in m.faculty for t in m.time) >= np.min([min_faculty, len(self.students_available[s]) ])

        # no student meets a faculty member more than once
        @m.Constraint(m.visitors, m.faculty)
        def meeting_each_prof_only_once(m, s, f):
            return sum(m.y[s, f, t] for t in m.time) <= 1

        # the number of students beyond one in each meeting
        @m.Constraint(m.faculty, m.time)
        def count_students_per_meeting(m, f, t):
            return m.beyond_one_visitor[f, t] >= sum(m.y[s, f, t] for s in m.visitors) - 1

        # number of faculty meeting more than max_visitors - 2 students
        @m.Constraint(m.faculty)
        def count_faculty_too_many_meetings(m, f):
            return 2*m.faculty_too_many_meetings[f] >= sum(m.y[s, f, t] for s in m.visitors for t in m.time) - max_visitors + 2

        @m.Constraint(m.faculty, m.time)
        def max_group_size(m, f, t):
            return sum(m.y[s, f, t] for s in m.visitors) <= max_group

        # Building phase offset: meetings cannot start before per-building phase slot.
        @m.Constraint(m.visitors, m.faculty, m.time)
        def building_phase(m, s, f, t):
            phase = self.building_phase_slot[self.faculty[f]["building"]]
            return m.y[s, f, t] <= (1 if t >= phase else 0)

        if self.movement_policy == MovementPolicy.TRAVEL_TIME.value:
            @m.Constraint(m.visitors, m.buildings, m.time)
            def in_building_link(m, s, b, t):
                return sum(m.y[s, f, t] for f in m.faculty_by_building[b]) <= m.in_building[s, b, t]

            m.travel_time_constraints = pyo.ConstraintList()
            for s in m.visitors:
                for b_from in self.buildings:
                    for b_to in self.buildings:
                        if b_from == b_to:
                            continue
                        lag = int(self.travel_slots[b_from][b_to])
                        if lag <= 0:
                            continue
                        for t_from in self.time_slots:
                            t_to_max = min(self.number_time_slots, t_from + lag)
                            for t_to in range(t_from + 1, t_to_max + 1):
                                m.travel_time_constraints.add(
                                    m.in_building[s, b_from, t_from] + m.in_building[s, b_to, t_to] <= 1
                                )

        if self.require_break_constraints_default or enforce_breaks:
            # Require at least one break for visitors within the configured break window
            @m.Constraint(m.visitors)
            def student_breaks(m, s):
                return sum(m.y[s, f, t] for f in m.faculty for t in m.break_options) <= len(m.break_options) - 1

        if self.require_break_constraints_default or enforce_breaks:
            # Determine when faculty are in breaks
            @m.Constraint(m.faculty, m.break_options)
            def faculty_in_break(m, f, t):
                return sum(m.y[s, f, t] for s in m.visitors) <= max_group*(1 - m.faculty_breaks[f,t])
            
            # Require at least one break for faculty
            @m.Constraint(m.faculty)
            def faculty_must_break(m, f):
                if len(self.faculty[f]["avail"]) < len(m.time):
                    # Faculty with limited availability do not get a break
                    return pyo.Constraint.Skip
                return sum(m.faculty_breaks[f,t] for t in m.break_options) >= 1

        for s in m.visitors:
            for t in m.time:
                if t not in self.students_available[s]:
                    for f in m.faculty:
                        m.y[s,f,t].fix(0)

        self.model = m
        
    def _solve_model(self, tee, record_results=True):
        """Solve the currently built model with the configured solver backend.

        Parameters
        ----------
        tee:
            If ``True``, stream solver output to stdout.
        """

        using_appsi_highs = False
        if self.solver is Solver.HIGHS:
            opt = None
            for solver_name in ("appsi_highs", "highs"):
                candidate = pyo.SolverFactory(solver_name)
                if candidate.available():
                    if solver_name == "appsi_highs":
                        using_appsi_highs = True
                        # Avoid RuntimeError on infeasible solves while still
                        # allowing explicit load for feasible solutions below.
                        if hasattr(candidate, "config") and hasattr(candidate.config, "load_solution"):
                            candidate.config.load_solution = False
                    opt = candidate
                    break
            if opt is None:
                raise RuntimeError("No available HiGHS solver found. Install appsi_highs or highs.")

        elif self.solver is Solver.CBC:
            opt = pyo.SolverFactory('cbc')
            # opt.options['ratio'] = 0.1 # set the ratio gap to 10%
            
        elif self.solver is Solver.GUROBI:
            opt = pyo.SolverFactory('gurobi')
            
        elif self.solver is Solver.GUROBI_IIS:
            
            # Reference: https://github.com/Pyomo/pyomo/blob/main/examples/pyomo/suffixes/gurobi_ampl_iis.py
            iis_exe = os.environ.get("GVS_GUROBI_IIS_EXECUTABLE", "./ampl/gurobi_ampl")
            iis_exe_path = Path(iis_exe)
            if not iis_exe_path.exists():
                raise RuntimeError(
                    f"GUROBI_IIS executable not found at '{iis_exe_path}'. "
                    "Set environment variable GVS_GUROBI_IIS_EXECUTABLE to the executable path."
                )
            opt = pyo.SolverFactory(str(iis_exe_path), solver_io='nl')
            if not opt.available(exception_flag=False):
                raise RuntimeError(
                    f"GUROBI_IIS solver is not available via '{iis_exe_path}'. "
                    "Check executable permissions and Gurobi/AMPL setup."
                )
            
            # tell gurobi to be verbose with output
            opt.options['outlev'] = 1
            
            # tell gurobi to find an iis table for the infeasible model
            opt.options['iisfind'] = 1  # tell gurobi to be verbose with output

            # Create an IMPORT Suffix to store the iis information that will be returned by gurobi_ampl
            self.model.iis = Suffix(direction=Suffix.IMPORT)
        
        try:
            results = opt.solve(self.model, tee=tee)
        except RuntimeError as exc:
            # Some HiGHS/Pyomo interfaces raise on infeasible models when
            # attempting to auto-load a non-existent solution.
            if "A feasible solution was not found" not in str(exc):
                raise
            results = opt.solve(self.model, tee=tee, load_solutions=False)
        if using_appsi_highs and results.solver.termination_condition in [TerminationCondition.optimal, TerminationCondition.feasible]:
            if hasattr(opt, "load_vars"):
                opt.load_vars()
        
        if self.solver is Solver.GUROBI_IIS:
            print("")
            print("IIS Results")
            for component, value in self.model.iis.items():
                print(component.name + " " + str(value))
        
        if record_results:
            self.results = results
            self.last_termination_condition = results.solver.termination_condition
            self.last_solver_status = results.solver.status
        return results

    def _build_solution_context(self):
        """Build immutable metadata context for solution snapshots."""
        times_by_building = {
            b: tuple(slots) for b, slots in self.times_by_building.items()
        }
        faculty = {}
        for name, info in self.faculty.items():
            faculty[name] = {
                "building": info.get("building", self.building_a),
                "room": info.get("room", ""),
                "avail": tuple(info.get("avail", [])),
                "areas": tuple(info.get("areas", [])),
            }
        requests = {s: tuple(v) for s, v in self.requests.items()}
        student_preferences = {
            (str(s), str(f)): float(v) for (s, f), v in self.student_preferences.items()
        }
        return SolutionContext(
            times_by_building=times_by_building,
            faculty=faculty,
            box_colors=dict(self.box_colors),
            number_time_slots=self.number_time_slots,
            run_name=self.run_name,
            student_preferences=student_preferences,
            requests=requests,
            legacy_faculty=frozenset(getattr(self, "legacy_faculty", {}).keys()),
            external_faculty=frozenset(self.external_faculty.keys()),
        )

    def _snapshot_solution(self, rank):
        """Capture the current solved model values as an immutable snapshot."""
        m = self.model
        visitors = tuple(str(s) for s in m.visitors)
        faculty = tuple(str(f) for f in m.faculty)
        time_slots = tuple(int(t) for t in m.time)
        active_meetings = frozenset(
            (str(s), str(f), int(t))
            for s in m.visitors
            for f in m.faculty
            for t in m.time
            if m.y[s, f, t]() >= 0.5
        )
        return SolutionResult(
            rank=rank,
            objective_value=float(pyo.value(m.obj)),
            termination_condition=str(self.last_termination_condition),
            solver_status=str(self.last_solver_status),
            visitors=visitors,
            faculty=faculty,
            time_slots=time_slots,
            active_meetings=active_meetings,
            context=self._build_solution_context(),
        )

    def _add_no_good_cut(self, solution):
        """Exclude an already-found binary assignment from future solves."""
        m = self.model
        if not hasattr(m, "no_good_cuts"):
            m.no_good_cuts = pyo.ConstraintList()

        all_indices = [(s, f, t) for s in m.visitors for f in m.faculty for t in m.time]
        active_indices = {
            (s, f, t)
            for s in m.visitors
            for f in m.faculty
            for t in m.time
            if solution.meeting_assigned(str(s), str(f), int(t))
        }
        m.no_good_cuts.add(
            sum(1 - m.y[s, f, t] for (s, f, t) in active_indices)
            + sum(m.y[s, f, t] for (s, f, t) in all_indices if (s, f, t) not in active_indices)
            >= 1
        )

    def has_feasible_solution(self):
        """Return ``True`` if the latest solver run was feasible or optimal."""
        if not hasattr(self, "results"):
            return False
        return self.last_termination_condition in [TerminationCondition.optimal, TerminationCondition.feasible]

    def infeasibility_report(self):
        """Summarize likely causes when the model has no feasible solution."""
        if not hasattr(self, "results"):
            return "No solver results available."
        if self.has_feasible_solution():
            return "Solution is feasible."
        params = getattr(self, "last_solve_params", {})
        min_visitors = params.get("min_visitors", 0)
        min_faculty = params.get("min_faculty", 0)
        max_group = params.get("max_group", 2)

        faculty_available = [f for f in self.faculty.keys() if len(self.faculty[f]["avail"]) > 0]
        num_faculty = len(faculty_available)
        num_students = len(self.student_data.index)
        total_faculty_capacity = sum(len(self.faculty[f]["avail"]) * max_group for f in faculty_available)
        total_student_capacity = sum(len(self.students_available[s]) for s in self.student_data.index)
        total_capacity = min(total_faculty_capacity, total_student_capacity)
        total_min_student_meetings = sum(min(min_faculty, len(self.students_available[s])) for s in self.student_data.index)
        total_min_faculty_meetings = num_faculty * min_visitors

        per_faculty_violations = []
        if min_visitors > 0:
            for f in faculty_available:
                max_meetings = len(self.faculty[f]["avail"]) * max_group
                if max_meetings < min_visitors:
                    per_faculty_violations.append((f, max_meetings))

        lines = [
            f"Termination: {self.last_termination_condition}, status: {self.last_solver_status}",
            f"Students: {num_students}, Faculty (available): {num_faculty}, Time slots: {self.number_time_slots}",
            f"Capacity (students): {total_student_capacity}, Capacity (faculty): {total_faculty_capacity}, Effective capacity: {total_capacity}",
            f"Required meetings from students (min_faculty): {total_min_student_meetings}",
            f"Required meetings from faculty (min_visitors): {total_min_faculty_meetings}",
        ]
        if total_min_student_meetings > total_capacity:
            lines.append("Infeasibility likely: min_faculty requirement exceeds total capacity.")
        if total_min_faculty_meetings > total_capacity:
            lines.append("Infeasibility likely: min_visitors requirement exceeds total capacity.")
        if per_faculty_violations:
            sample = ", ".join([f"{f} (max {m})" for f, m in per_faculty_violations[:6]])
            lines.append(f"Faculty with insufficient availability for min_visitors: {sample}")
        lines.append("Try reducing min_faculty/min_visitors or loosening availability.")
        return "\n".join(lines)

    def _current_solution_result(self):
        """Return a rich snapshot for the currently loaded feasible model."""
        if not self.has_feasible_solution():
            raise RuntimeError(
                f"No feasible solution available (termination: {getattr(self, 'last_termination_condition', None)})."
            )
        return self._snapshot_solution(rank=1)

    def show_faculty_schedule(self, save_files=True, abbreviate_student_names=True, solution=None, show_solution_rank=False):
        """Plot the solved schedule grouped by faculty.

        Parameters
        ----------
        save_files:
            If ``True``, save ``faculty_schedule*.pdf`` and ``.png``.
        abbreviate_student_names:
            If ``True``, shorten visitor names in labels.
        solution:
            Optional :class:`SolutionResult` to plot. If omitted, uses the most
            recently solved model on this scheduler.
        show_solution_rank:
            If ``True`` and ``solution`` is provided, append solution rank to
            the figure title.
        """
        warnings.warn(
            "Scheduler.show_faculty_schedule(...) is a legacy interface. "
            "Prefer SolutionResult.plot_faculty_schedule(...) via "
            "Scheduler.last_solution_set.get(rank) or SolutionSet.plot_faculty_schedule(...).",
            FutureWarning,
            stacklevel=2,
        )
        from_scheduler_state = solution is None
        chosen = solution if solution is not None else self._current_solution_result()
        return chosen.plot_faculty_schedule(
            save_files=save_files,
            abbreviate_student_names=abbreviate_student_names,
            show_solution_rank=show_solution_rank,
            include_rank_in_filename=not from_scheduler_state,
        )
    
    def show_visitor_schedule(self, save_files=True, abbreviate_student_names=True, solution=None, show_solution_rank=False):
        """Plot the solved schedule grouped by visitor.

        Parameters
        ----------
        save_files:
            If ``True``, save ``visitor_schedule*.pdf`` and ``.png``.
        abbreviate_student_names:
            If ``True``, shorten visitor names in y-axis labels.
        solution:
            Optional :class:`SolutionResult` to plot. If omitted, uses the most
            recently solved model on this scheduler.
        show_solution_rank:
            If ``True`` and ``solution`` is provided, append solution rank to
            the figure title.
        """
        warnings.warn(
            "Scheduler.show_visitor_schedule(...) is a legacy interface. "
            "Prefer SolutionResult.plot_visitor_schedule(...) via "
            "Scheduler.last_solution_set.get(rank) or SolutionSet.plot_visitor_schedule(...).",
            FutureWarning,
            stacklevel=2,
        )
        from_scheduler_state = solution is None
        chosen = solution if solution is not None else self._current_solution_result()
        return chosen.plot_visitor_schedule(
            save_files=save_files,
            abbreviate_student_names=abbreviate_student_names,
            show_solution_rank=show_solution_rank,
            include_rank_in_filename=not from_scheduler_state,
        )

    def export_visitor_docx(self, filename, solution=None, **kwargs):
        """Export solved visitor schedules to a DOCX file.

        Parameters
        ----------
        filename:
            Destination DOCX path.
        **kwargs:
            Additional keyword arguments forwarded to
            :func:`grad_visit_scheduler.export.export_visitor_docx`.
        """
        warnings.warn(
            "Scheduler.export_visitor_docx(...) is a legacy interface. "
            "Prefer SolutionResult.export_visitor_docx(...) via "
            "Scheduler.last_solution_set.get(rank) or SolutionSet.export_visitor_docx(...).",
            FutureWarning,
            stacklevel=2,
        )
        chosen = solution if solution is not None else self._current_solution_result()
        return chosen.export_visitor_docx(filename, **kwargs)
        
    def show_utility(self):
        """Plot realized meetings and display total utility for a solved model."""
        m = self.model
        y = pd.DataFrame.from_dict({f:{s: sum(m.y[s, f, t]() for t in m.time) for s in m.visitors} for f in m.faculty})
        fig, ax = plt.subplots(1, 1, figsize=(10, 8))
        sns.heatmap(y, ax=ax, annot=True, cmap="crest", square=True, cbar=False)
        utility = sum(m.utility[s]() for s in m.visitors)
        ax.set_title(f"Total Utility = {utility:0.1f}")
        return fig, ax

    def check_requests(self):
        """Print unmet requested faculty meetings and meeting-count distribution."""
        if not self.has_feasible_solution():
            raise RuntimeError(self.infeasibility_report())
        m = self.model
        total_meetings = {}
        for s in m.visitors:
            total_meetings[s] = 0
            for f in m.faculty:
                preferred_meeting = 0
                for t in m.time:
                    total_meetings[s] += int(m.y[s, f, t]())
                
                    if f in self.requests[s]:
                        preferred_meeting += m.y[s, f, t]()
                if f in self.requests[s] and preferred_meeting < 0.99:
                    print(s,"is not meeting with",f, "( weight =",self.student_preferences[(s, f)],")")
        
        print(" ")
        num_meetings = [v for i, (k, v) in enumerate(total_meetings.items())]
        count = Counter(num_meetings)
        for i, k in enumerate(count):
            print(count[k],"visitors with",k,"meetings")
        
        #for s in total_meetings.keys():
        #    print(s,"has",total_meetings[s],"meetings")
